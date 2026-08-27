"""정책 문서에서 Chroma 인덱스를 부트스트랩한다 (배포용, LlamaCloud 불필요)."""

from __future__ import annotations

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config.config import DOCX_DATA_PATH


def load_policy_documents_from_docx() -> list[Document]:
    """docx를 단락 단위로 읽어 Document 목록을 만든다."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError(
            "python-docx가 필요합니다. pip install python-docx"
        ) from exc

    if not DOCX_DATA_PATH.is_file():
        raise FileNotFoundError(f"정책 문서가 없습니다: {DOCX_DATA_PATH}")

    doc = DocxDocument(str(DOCX_DATA_PATH))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        raise ValueError(f"정책 문서에 텍스트가 없습니다: {DOCX_DATA_PATH}")

    full_text = "\n\n".join(paragraphs)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=120,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(full_text)
    return [
        Document(
            page_content=chunk,
            metadata={"source": DOCX_DATA_PATH.name, "chunk_index": index},
        )
        for index, chunk in enumerate(chunks)
    ]
